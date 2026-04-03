"""Extract plain text from PDF and text-based files (path or in-memory bytes)."""
from __future__ import annotations

import hashlib
import io
from pathlib import Path

from django.conf import settings


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_text_from_bytes(data: bytes, suffix: str) -> str:
    """Parse upload bytes in memory (no temp file). suffix includes leading dot, e.g. .docx."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        return _pdf_text_bytes(data)
    if suffix in {".txt", ".md", ".markdown"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return _docx_text_bytes(data)
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, .md, or .docx")


def _docx_text_bytes(data: bytes) -> str:
    import docx

    buf = io.BytesIO(data)
    document = docx.Document(buf)
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _pdf_text_bytes(data: bytes) -> str:
    from pypdf import PdfReader

    buf = io.BytesIO(data)
    reader = PdfReader(buf)
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip()


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _docx_text(path)
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, .md, or .docx")


def _docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip()


def truncate(text: str, max_chars: int | None = None) -> str:
    max_chars = max_chars or getattr(settings, "SYNDICATE_MAX_DOC_CHARS", 120_000)
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[... document truncated for processing ...]"
